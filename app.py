from flask import Flask, render_template, request, redirect, url_for, session, jsonify, Response
from functools import wraps
from urllib.parse import urlparse, urljoin
import json
from datetime import datetime, timedelta
import pandas as pd
import pytz
from pymongo import MongoClient
import pymongo
from bson import ObjectId
import os
import logging
import re
import requests  # ✅ ADD THIS IMPORT
from werkzeug.utils import secure_filename  # ✅ ADD THIS IMPORT
import mimetypes
import io
try:
    import openai
except ImportError:
    openai = None
from dotenv import load_dotenv

try:
    from openai import OpenAI as OpenAIClient
except ImportError:
    OpenAIClient = None

# Text extraction libraries
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

load_dotenv()

LOG_LEVEL_NAME = os.getenv('APP_LOG_LEVEL', 'INFO').upper()
LOG_LEVEL = getattr(logging, LOG_LEVEL_NAME, logging.INFO)
logging.basicConfig(
    level=LOG_LEVEL,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s'
)
logger = logging.getLogger('antoine_dashboard')

app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'change-me-in-env')
app.permanent_session_lifetime = timedelta(days=1)

MONGODB_URI = os.getenv(
    'MONGODB_URI',
    "mongodb+srv://userawais1:awais645@cluster0.fcqph.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"
)
MONGODB_DB_NAME = os.getenv('MONGODB_DB_NAME', 'chat_history')
DEFAULT_MONGO_COLLECTION = os.getenv('MONGODB_DEFAULT_COLLECTION', 'VincentBotReplicaWithVoice')

OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
OPENAI_SUMMARY_MODEL = os.getenv('OPENAI_SUMMARY_MODEL', 'gpt-4o-mini')

if OPENAI_API_KEY and (openai is not None or OpenAIClient is not None):
    if OpenAIClient:
        _openai_client = OpenAIClient(api_key=OPENAI_API_KEY)
    else:
        _openai_client = None
        if openai:
            openai.api_key = OPENAI_API_KEY
else:
    _openai_client = None

SUMMARY_PLACEHOLDER = "Summary not available."
TEXT_SUMMARY_EXTENSIONS = {
    'txt', 'md', 'csv', 'json', 'xml', 'py', 'js', 'ts', 'tsx', 'html', 'htm',
    'css', 'log'
}
MAX_SUMMARY_INPUT_CHARS = 6000

# Azure backend base URL - Global configuration
BASE_URL = "https://capps-backend-ooxig22w5exq6.calmpebble-4f694259.westus.azurecontainerapps.io"


def is_safe_url(target: str) -> bool:
    """
    Ensure the target URL is safe to redirect to (same host and http/https scheme)
    """
    if not target:
        return False
    ref_url = urlparse(request.host_url)
    test_url = urlparse(urljoin(request.host_url, target))
    return (
        test_url.scheme in ("http", "https")
        and ref_url.netloc == test_url.netloc
    )


def login_required(view_func):
    """
    Decorator to enforce authentication for view functions.
    Redirects unauthenticated users to the login page and preserves the
    original destination (when safe) using the `next` query parameter.
    """
    @wraps(view_func)
    def wrapped_view(*args, **kwargs):
        if 'email' not in session:
            next_url = request.full_path if request.method in ('GET', 'HEAD') else ''
            if not is_safe_url(next_url):
                next_url = ''
            return redirect(url_for('login', next=next_url or None))
        return view_func(*args, **kwargs)
    return wrapped_view

# MongoDB connection setup
_mongo_client = None


def get_mongo_client():
    global _mongo_client
    if _mongo_client is None:
        if not MONGODB_URI:
            raise RuntimeError("MONGODB_URI environment variable is not set.")
        _mongo_client = MongoClient(MONGODB_URI)
    return _mongo_client


def get_mongo_connection(collection_name=DEFAULT_MONGO_COLLECTION):
    client = get_mongo_client()
    db = client[MONGODB_DB_NAME]
    collection = db[collection_name]
    return collection

def read_chat_data(user_email=None, user_role=None):
    try:
        print("Reading data from MongoDB collection...")
        
        # Get data from VincentBotReplicaWithVoice collection only
        collection = get_mongo_connection()
        
        # Create query filter based on user email and role
        query_filter = {}
        if user_role == 'member' and user_email:
            # If user is member, only show their own data
            query_filter['email'] = user_email
            print(f"Member user - Filtering data for user: {user_email}")
        elif user_role == 'admin':
            # If user is admin, show all data (no filter)
            print("Admin user - Showing all data")
        else:
            # Default to member behavior if role is not specified
            query_filter['email'] = user_email
            print(f"Default behavior - Filtering data for user: {user_email}")
        
        cursor = collection.find(query_filter).sort("timestamp", pymongo.DESCENDING)
        df = pd.DataFrame(list(cursor))
        print(f"Retrieved {len(df)} records from collection")
        
        # Check if dataframe is empty
        if df.empty:
            print("No data found in MongoDB collection")
            return None
        
        # Field mapping based on your document structure
        field_mapping = {
            'email': 'Email',
            'user_query': 'Question',      # ✅ This is correct
            'bot_response': 'Answer', 
            'timestamp': 'Date',
            'feedback': 'Feedback'
        }
        
        # Apply field mappings
        for mongo_field, expected_field in field_mapping.items():
            if mongo_field in df.columns:
                df[expected_field] = df[mongo_field]
                print(f"Mapped {mongo_field} to {expected_field}")
        
        # Handle timestamp field (keep your existing timestamp processing code)
        if 'timestamp' in df.columns:
            print("Converting timestamp to datetime...")
            
            def convert_timestamp(ts):
                if isinstance(ts, dict) and '$date' in ts:
                    # Handle MongoDB $date format
                    if '$numberLong' in ts['$date']:
                        ts_long = int(ts['$date']['$numberLong'])
                        return datetime.fromtimestamp(ts_long / 1000, tz=pytz.UTC)
                    else:
                        # Handle other date formats if needed
                        return pd.NaT
                else:
                    # Try to parse as is
                    try:
                        return pd.to_datetime(ts)
                    except:
                        return pd.NaT
            
            # Apply the conversion function
            df['Date'] = df['timestamp'].apply(convert_timestamp)
            print(f"Sample converted dates: {df['Date'].head(3)}")
        elif 'Date' not in df.columns:
            # Use ObjectId generation time as fallback
            print("No timestamp found, using _id to generate dates...")
            df['Date'] = df['_id'].apply(lambda x: x.generation_time if isinstance(x, ObjectId) else pd.NaT)
        
        # Make sure Date column is timezone aware
        if 'Date' in df.columns and not df['Date'].empty:
            # Ensure all dates are timezone aware
            if not pd.api.types.is_datetime64tz_dtype(df['Date']):
                df['Date'] = df['Date'].dt.tz_localize(pytz.UTC)
            
            print(f"Date column dtype: {df['Date'].dtype}")
            print(f"Sample dates: {df['Date'].head(3)}")
        
        # Sort by date
        df = df.sort_values(by='Date', ascending=False)
        
        return df
    except Exception as e:
        print(f"Error reading MongoDB data: {e}")
        import traceback
        traceback.print_exc()
        return None

def get_data_by_period(period='today', start_date=None, end_date=None, user_filter=None, user_email=None, user_role=None):
    df = read_chat_data(user_email=user_email, user_role=user_role)
    if df is None:
        # Return empty data structure
        return {
            "labels": [],
            "values": [],
            "total_messages": 0,
            "tickets_created": 0,
            "user_percentages": [],
            "feedback_stats": {"Good": 0, "Bad": 0, "Neutral": 0}
        }
    
    # Make sure we're using timezone-aware datetime objects for comparison
    current_time = datetime.now(pytz.UTC)
    
    # Filter data based on selected period
    if period == 'custom' and start_date and end_date:
        try:
            start_date = datetime.strptime(start_date, "%Y-%m-%d").replace(hour=0, minute=0, second=0, microsecond=0)
            end_date = datetime.strptime(end_date, "%Y-%m-%d").replace(hour=23, minute=59, second=59, microsecond=999)
            # Make timezone aware
            start_date = pytz.UTC.localize(start_date)
            end_date = pytz.UTC.localize(end_date)
            df_filtered = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]
        except ValueError:
            # If date parsing fails, fallback to today
            start_date = current_time.replace(hour=0, minute=0, second=0, microsecond=0)
            end_date = current_time
            df_filtered = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]
    elif period == 'today':
        start_date = current_time.replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = current_time
        df_filtered = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]
    elif period == 'yesterday':
        start_date = (current_time - timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = start_date + timedelta(days=1) - timedelta(microseconds=1)
        df_filtered = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]
    elif period == 'this-week':
        start_date = (current_time - timedelta(days=current_time.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
        df_filtered = df[df['Date'] >= start_date]
    elif period == 'last-week':
        start_date = (current_time - timedelta(days=current_time.weekday() + 7)).replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = start_date + timedelta(days=7) - timedelta(microseconds=1)
        df_filtered = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]
    elif period == 'this-month':
        start_date = current_time.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        df_filtered = df[df['Date'] >= start_date]
    elif period == 'last-month':
        last_month = current_time.month - 1 if current_time.month > 1 else 12
        last_month_year = current_time.year if current_time.month > 1 else current_time.year - 1
        start_date = current_time.replace(year=last_month_year, month=last_month, day=1, hour=0, minute=0, second=0, microsecond=0)
        end_date = current_time.replace(day=1, hour=0, minute=0, second=0, microsecond=0) - timedelta(microseconds=1)
        df_filtered = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]
    else:
        # Default to all data
        df_filtered = df
    
    # Apply user filter if provided (additional filtering for admin users)
    if user_filter and isinstance(user_filter, list) and len(user_filter) > 0 and 'Email' in df_filtered.columns:
        df_filtered = df_filtered[df_filtered['Email'].isin(user_filter)]
        print(f"Applied additional user filter: {user_filter}. Remaining messages: {len(df_filtered)}")
    
    # Calculate feedback statistics
    feedback_stats = {"Good": 0, "Bad": 0, "Neutral": 0}
    if 'Feedback' in df_filtered.columns:
        feedback_counts = df_filtered['Feedback'].value_counts()
        for feedback, count in feedback_counts.items():
            if feedback in feedback_stats:
                feedback_stats[feedback] = int(count)
            elif feedback == "":
                feedback_stats["Neutral"] += int(count)
    
    # Group by appropriate time interval based on the period
    if period in ['today', 'yesterday']:
        # Group by hour for single day views
        df_filtered['Hour'] = df_filtered['Date'].dt.hour
        df_filtered['Minute'] = df_filtered['Date'].dt.minute
        
        if not df_filtered.empty:
            # Get the latest time from the filtered data
            latest_record = df_filtered.loc[df_filtered['Date'].idxmax()]
            latest_hour = latest_record['Hour']
            latest_minute = latest_record['Minute']
            latest_time_str = f"{int(latest_hour):02d}:{int(latest_minute):02d}"
            
            # Group by hour
            hourly_counts = df_filtered.groupby('Hour').size().reset_index(name='count')
            
            # Generate a complete set of hours (0-latest_hour) with counts
            all_hours = pd.DataFrame({'Hour': range(latest_hour + 1)})
            hourly_counts = pd.merge(all_hours, hourly_counts, on='Hour', how='left').fillna(0)
            
            # Format labels as HH:00, with the last one showing the exact time
            labels = [f"{int(hour):02d}:00" for hour in hourly_counts['Hour']]
            if labels:
                labels[-1] = latest_time_str  # Replace the last label with the exact time
        else:
            # If no data, show empty hours
            labels = []
            values = []
            return {
                "labels": labels,
                "values": values,
                "total_messages": 0,
                "tickets_created": 0,
                "user_percentages": [],
                "feedback_stats": feedback_stats
            }
        
        values = hourly_counts['count'].astype(int).tolist()
    else:
        # Group by date for multi-day views
        df_filtered['Date_Day'] = df_filtered['Date'].dt.date
        
        if not df_filtered.empty:
            daily_counts = df_filtered.groupby('Date_Day').size().reset_index(name='count')
        
            # Prepare data for chart
            labels = [str(date) for date in daily_counts['Date_Day']]
            values = daily_counts['count'].tolist()
        else:
            # If no data, return empty arrays
            labels = []
            values = []
            return {
                "labels": labels,
                "values": values,
                "total_messages": 0,
                "tickets_created": 0,
                "user_percentages": [],
                "feedback_stats": feedback_stats
            }
    
    # Calculate metrics
    total_messages = len(df_filtered)
    # Count unique emails that created tickets (assume a message creates a ticket)
    tickets_created = df_filtered['Email'].nunique() if 'Email' in df_filtered.columns else 0
    
    # Calculate user chat percentages for donut chart
    user_percentages = []
    if not df_filtered.empty and 'Email' in df_filtered.columns:
        try:
            print("Calculating user percentages...")
            # Clean data
            df_filtered_clean = df_filtered.dropna(subset=['Email'])
            df_filtered_clean = df_filtered_clean[df_filtered_clean['Email'].str.strip() != '']
            
            if len(df_filtered_clean) > 0:
                user_counts = df_filtered_clean.groupby('Email').size().reset_index(name='count')
                print("User counts:", user_counts)
                
                total_messages = len(df_filtered_clean)
                print(f"Total messages: {total_messages}")
                
                if total_messages > 0:
                    user_counts['percentage'] = (user_counts['count'] / total_messages * 100).round(1)
                    user_counts = user_counts.sort_values('count', ascending=False)
                    
                    # Create list of user data for the donut chart
                    for _, row in user_counts.iterrows():
                        email = row['Email']
                        if email and isinstance(email, str):
                            user_percentages.append({
                                "username": email,
                                "count": int(row['count']),
                                "percentage": float(row['percentage'])
                            })
        except Exception as e:
            print(f"Error calculating user percentages: {e}")
            import traceback
            traceback.print_exc()
    
    # Prepare data for response
    chat_data = {
        "labels": labels,
        "values": values,
        "total_messages": total_messages,
        "tickets_created": tickets_created,
        "user_percentages": user_percentages,
        "feedback_stats": feedback_stats
    }
    
    print(f"Final response data structure: {chat_data}")
    return chat_data

@app.route('/')
def dashboard():
    if 'email' not in session:
        return redirect(url_for('login'))
    return render_template('index.html')

# Login page
@app.route('/login', methods=['GET', 'POST'])
def login():
    next_url = request.args.get('next', '')

    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        next_url = request.form.get('next', next_url)

        users_collection = get_mongo_connection('users')
        user = users_collection.find_one({'email': email, 'password': password})

        if user:
            session.permanent = True  
            session['email'] = email
            session['role'] = user.get('role', 'member')  # Get role from database, default to 'member'
            print(f"User {email} logged in with role: {session['role']}")
            if next_url and is_safe_url(next_url):
                return redirect(next_url)
            return redirect(url_for('dashboard'))
        else:
            return render_template('login.html', error='Invalid credentials', next=next_url)

    return render_template('login.html', next=next_url)

@app.route('/logout')
def logout():
    session.pop('email', None)  # Remove user from session
    session.pop('role', None)   # Remove role from session
    return redirect(url_for('login'))  # Redirect to login page

@app.route('/api/chat_data')
def get_chat_data():
    # Check if user is authenticated
    if 'email' not in session:
        return jsonify({"error": "Unauthorized"}), 401
    
    period = request.args.get('period', 'today')
    start_date = request.args.get('start_date', None)
    end_date = request.args.get('end_date', None)
    
    # Get user filter if provided
    user_filter = None
    users_param = request.args.get('users', None)
    if users_param:
        try:
            user_filter = json.loads(users_param)
            print(f"Received user filter: {user_filter}")
        except json.JSONDecodeError:
            print(f"Invalid JSON in users parameter: {users_param}")
    
    # Pass the authenticated user's email and role to filter data
    authenticated_email = session['email']
    user_role = session.get('role', 'member')
    
    data = get_data_by_period(period, start_date, end_date, user_filter, user_email=authenticated_email, user_role=user_role)
    return jsonify(data)

@app.route('/data')
def get_filtered_data():
    # Check if user is authenticated
    if 'email' not in session:
        return jsonify({"error": "Unauthorized"}), 401
    
    start_date = request.args.get('start', None)
    end_date = request.args.get('end', None)
    
    # Get user filter if provided
    user_filter = None
    users_param = request.args.get('users', None)
    if users_param:
        try:
            user_filter = json.loads(users_param)
            print(f"Received user filter for /data route: {user_filter}")
        except json.JSONDecodeError:
            print(f"Invalid JSON in users parameter: {users_param}")
    
    if not start_date or not end_date:
        # If no date range is provided, use today
        current_time = datetime.now()
        start_date = current_time.strftime("%Y-%m-%d")
        end_date = current_time.strftime("%Y-%m-%d")
    
    # Get the filtered data with authenticated user's email and role
    authenticated_email = session['email']
    user_role = session.get('role', 'member')
    
    data = get_data_by_period('custom', start_date, end_date, user_filter, user_email=authenticated_email, user_role=user_role)
    
    # Format the response in the structure expected by the frontend
    response = {
        "linechart_data": {
            "labels": data["labels"],
            "values": data["values"]
        },
        "total_messages": data["total_messages"],
        "active_users": len(data["user_percentages"]) if data["user_percentages"] else 0,
        "user_data": data["user_percentages"],
        "feedback_stats": data["feedback_stats"],
        "period": "custom"
    }
    
    return jsonify(response)

@app.route('/chats')
@login_required
def chats():
    return render_template('chats.html')

@app.route('/api/chats')
def get_chat_list():
    # Check if user is authenticated
    if 'email' not in session:
        return jsonify({"error": "Unauthorized"}), 401
    
    # Get authenticated user's email and role
    authenticated_email = session['email']
    user_role = session.get('role', 'member')
    print(f"Authenticated user: {authenticated_email}, Role: {user_role}")
    df = read_chat_data(user_email=authenticated_email, user_role=user_role)
    if df is None:
        return jsonify({
            "chats": [],
            "total_count": 0
        })
    
    # Get filter parameters
    period = request.args.get('period')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    username = request.args.get('username')
    users_param = request.args.get('users')
    feedback = request.args.get('feedback')
    
    # Debug: Print the filter parameters received
    print(f"Filter parameters: period={period}, start_date={start_date}, end_date={end_date}, users={users_param}")
    print(f"User role: {user_role}")
    
    # Apply date filters - using timezone-aware dates
    current_time = datetime.now(pytz.UTC)
    
    # Apply proper date filtering based on period
    if period:
        print(f"Applying period filter: {period}")
        if period == 'today':
            start_date_obj = current_time.replace(hour=0, minute=0, second=0, microsecond=0)
            end_date_obj = current_time
            df = df[(df['Date'] >= start_date_obj) & (df['Date'] <= end_date_obj)]
            print(f"Today filter: {start_date_obj} to {end_date_obj}, records: {len(df)}")
        elif period == 'yesterday':
            start_date_obj = (current_time - timedelta(days=1)).replace(hour=0, minute=0, second=0, microsecond=0)
            end_date_obj = start_date_obj + timedelta(days=1) - timedelta(microseconds=1)
            df = df[(df['Date'] >= start_date_obj) & (df['Date'] <= end_date_obj)]
            print(f"Yesterday filter: {start_date_obj} to {end_date_obj}, records: {len(df)}")
        elif period == 'this-week':
            start_date_obj = (current_time - timedelta(days=current_time.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
            df = df[df['Date'] >= start_date_obj]
            print(f"This week filter: from {start_date_obj}, records: {len(df)}")
        elif period == 'last-week':
            start_date_obj = (current_time - timedelta(days=current_time.weekday() + 7)).replace(hour=0, minute=0, second=0, microsecond=0)
            end_date_obj = start_date_obj + timedelta(days=7) - timedelta(microseconds=1)
            df = df[(df['Date'] >= start_date_obj) & (df['Date'] <= end_date_obj)]
            print(f"Last week filter: {start_date_obj} to {end_date_obj}, records: {len(df)}")
        elif period == 'this-month':
            start_date_obj = current_time.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            df = df[df['Date'] >= start_date_obj]
            print(f"This month filter: from {start_date_obj}, records: {len(df)}")
        elif period == 'last-month':
            last_month = current_time.month - 1 if current_time.month > 1 else 12
            last_month_year = current_time.year if current_time.month > 1 else current_time.year - 1
            start_date_obj = current_time.replace(year=last_month_year, month=last_month, day=1, hour=0, minute=0, second=0, microsecond=0)
            end_date_obj = current_time.replace(day=1, hour=0, minute=0, second=0, microsecond=0) - timedelta(microseconds=1)
            df = df[(df['Date'] >= start_date_obj) & (df['Date'] <= end_date_obj)]
            print(f"Last month filter: {start_date_obj} to {end_date_obj}, records: {len(df)}")
    elif start_date and end_date:
        try:
            start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").replace(hour=0, minute=0, second=0, microsecond=0)
            end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").replace(hour=23, minute=59, second=59, microsecond=999)
            # Make timezone aware
            start_date_obj = pytz.UTC.localize(start_date_obj)
            end_date_obj = pytz.UTC.localize(end_date_obj)
            df = df[(df['Date'] >= start_date_obj) & (df['Date'] <= end_date_obj)]
            print(f"Custom date filter: {start_date_obj} to {end_date_obj}, records: {len(df)}")
        except ValueError as e:
            print(f"Date parsing error: {e}")
            # If date parsing fails, don't apply date filter

    # Apply username filter (only for admin users)
    if user_role == 'admin' and users_param:
        try:
            # Try to parse as JSON
            users = json.loads(users_param)
            print(f"Admin user - Parsed users_param: {users}")
            
            # Handle various formats that might come from frontend
            if isinstance(users, list):
                # Flatten nested lists
                flat_users = []
                for item in users:
                    if isinstance(item, list):
                        flat_users.extend([u for u in item if u and not pd.isna(u)])
                    elif item and not pd.isna(item):
                        flat_users.append(item)
                
                # Apply the filter if we have valid users
                if flat_users and 'Email' in df.columns:
                    df = df[df['Email'].isin(flat_users)]
                    print(f"Applied admin user filter with {len(flat_users)} users, remaining records: {len(df)}")
            elif isinstance(users, str) and users and 'Email' in df.columns:
                df = df[df['Email'] == users]
                print(f"Applied single admin user filter: {users}, remaining records: {len(df)}")
        except json.JSONDecodeError as e:
            print(f"Error parsing users JSON: {e}")
    elif user_role == 'admin' and username and 'Email' in df.columns:
        # Fallback to direct username filter for admin
        df = df[df['Email'].str.contains(str(username), case=False, na=False)]
        print(f"Applied admin username text search: {username}, remaining records: {len(df)}")
    
    # Apply feedback filter if provided
    if feedback and 'Feedback' in df.columns:
        df = df[df['Feedback'] == feedback]
        print(f"Applied feedback filter: {feedback}, remaining records: {len(df)}")
    
    # Sort by date (most recent first)
    df = df.sort_values(by='Date', ascending=False)
    
    # Format data for JSON response
    chats = []
    for _, row in df.iterrows():
        # Convert ObjectId to string for JSON serialization
        id_value = str(row.get('_id')) if row.get('_id') else str(_)
        
        # Extract conversation history if available
        conversation_history = []
        if 'message_history' in row and isinstance(row['message_history'], list):
            conversation_history = row['message_history']
        
        # Handle NaN values by replacing them with None (which becomes null in JSON)
        feedback_value = row.get('Feedback', "")
        if pd.isna(feedback_value):
            feedback_value = ""
        
        # Similarly handle other fields that might contain NaN
        email_value = row.get('Email', "Unknown")
        if pd.isna(email_value):
            email_value = "Unknown"
            
        question_value = row.get('Question', "")
        if pd.isna(question_value):
            question_value = ""
            
        answer_value = row.get('Answer', "")
        if pd.isna(answer_value):
            answer_value = ""
            
        conv_id = row.get('conversation_id', "")
        if pd.isna(conv_id):
            conv_id = ""
        
        chat = {
            "id": id_value,
            "date": row['Date'].strftime("%d-%b-%Y %H:%M") if 'Date' in row and not pd.isna(row['Date']) else "",
            "bot_type": "ChatBot",
            "username": email_value,
            "question": question_value,
            "answer": answer_value,
            "feedback": feedback_value,
            "conversation_id": conv_id,
            "conversation_history": conversation_history
        }
        chats.append(chat)
    
    # Print summary for debugging
    print(f"Final chat count after all filters: {len(chats)}")
    
    return jsonify({
        "chats": chats,
        "total_count": len(chats)
    })

# Add API endpoint for username suggestions - role-based filtering
@app.route('/api/usernames')
def get_usernames():
    # Check if user is authenticated
    if 'email' not in session:
        return jsonify({"error": "Unauthorized"}), 401
    
    try:
        authenticated_email = session['email']
        user_role = session.get('role', 'member')
        collection = get_mongo_connection()
        
        if user_role == 'admin':
            # Admin can see all usernames
            unique_emails = collection.distinct('email')
            # Filter out empty values and sort
            unique_emails = [email for email in unique_emails if email and isinstance(email, str) and email.strip()]
            unique_emails.sort()
        else:
            # Member can only see their own email
            unique_emails = [authenticated_email]
        
        return jsonify({
            "usernames": unique_emails
        })
    except Exception as e:
        print(f"Error retrieving usernames from MongoDB: {e}")
        return jsonify({
            "usernames": []
        })

# Add API endpoint for feedback filters - role-based filtering
@app.route('/api/feedback_options')
def get_feedback_options():
    # Check if user is authenticated
    if 'email' not in session:
        return jsonify({"error": "Unauthorized"}), 401
    
    try:
        authenticated_email = session['email']
        user_role = session.get('role', 'member')
        collection = get_mongo_connection()
        
        if user_role == 'admin':
            # Admin can see feedback from all users
            unique_feedback = collection.distinct('feedback')
        else:
            # Member can only see their own feedback options
            unique_feedback = collection.distinct('feedback', {'email': authenticated_email})
        
        # Filter out empty values and sort
        unique_feedback = [feedback for feedback in unique_feedback if feedback and isinstance(feedback, str) and feedback.strip()]
        unique_feedback.sort()
        
        return jsonify({
            "feedback_options": unique_feedback
        })
    except Exception as e:
        print(f"Error retrieving feedback options from MongoDB: {e}")
        return jsonify({
            "feedback_options": ["Good", "Bad", "Neutral"]  # Default options
        })

# Add API endpoint to get user info (useful for frontend to know user role)
@app.route('/api/user_info')
def get_user_info():
    # Check if user is authenticated
    if 'email' not in session:
        return jsonify({"error": "Unauthorized"}), 401
    
    return jsonify({
        "email": session['email'],
        "role": session.get('role', 'member'),
        "is_admin": session.get('role', 'member') == 'admin'
    })
    
@app.route('/upload')
@login_required
def upload_file():
    return render_template('upload.html')



# Upload configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {
    'jpg', 'jpeg', 'png', 'pdf', 'doc', 'docx', 'txt', 'xlsx', 'csv', 'gif', 'zip', 'rar', 'json', 'xlsb'
}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB in bytes

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max request size

# Create upload folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def is_text_like_file(filename, content_type=None):
    if not filename:
        return False
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if ext in TEXT_SUMMARY_EXTENSIONS:
        return True
    if content_type:
        content_type = content_type.lower()
        if content_type.startswith('text/'):
            return True
        if content_type in ('application/json', 'application/xml'):
            return True
    return False


def extract_text_from_file(filename, file_bytes, content_type=None):
    """
    Extract text from various file types (PDF, DOCX, TXT, CSV, Excel, etc.)
    Returns extracted text or empty string if extraction fails
    """
    if not file_bytes or not filename:
        return ''
    
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    
    try:
        # PDF files
        if ext == 'pdf' and PyPDF2:
            try:
                pdf_file = io.BytesIO(file_bytes)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text_parts = []
                # Limit to first 50 pages to avoid processing huge files
                max_pages = min(50, len(pdf_reader.pages))
                for page_num in range(max_pages):
                    page = pdf_reader.pages[page_num]
                    text_parts.append(page.extract_text())
                extracted_text = '\n'.join(text_parts)
                if extracted_text.strip():
                    print(f"✅ Extracted {len(extracted_text)} characters from PDF: {filename}")
                    return extracted_text.strip()
            except Exception as e:
                print(f"⚠️ PDF extraction failed for {filename}: {e}")
        
        # DOCX files
        elif ext == 'docx' and DocxDocument:
            try:
                docx_file = io.BytesIO(file_bytes)
                doc = DocxDocument(docx_file)
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)
                extracted_text = '\n'.join(text_parts)
                if extracted_text.strip():
                    print(f"✅ Extracted {len(extracted_text)} characters from DOCX: {filename}")
                    return extracted_text.strip()
            except Exception as e:
                print(f"⚠️ DOCX extraction failed for {filename}: {e}")
        
        # Excel files (XLSX, XLSB)
        elif ext in ('xlsx', 'xlsb'):
            try:
                excel_file = io.BytesIO(file_bytes)
                # Read all sheets
                df = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
                text_parts = []
                for sheet_name, sheet_df in df.items():
                    # Convert DataFrame to text representation
                    text_parts.append(f"Sheet: {sheet_name}")
                    text_parts.append(sheet_df.to_string())
                extracted_text = '\n'.join(text_parts)
                if extracted_text.strip():
                    print(f"✅ Extracted {len(extracted_text)} characters from Excel: {filename}")
                    return extracted_text.strip()
            except Exception as e:
                print(f"⚠️ Excel extraction failed for {filename}: {e}")
        
        # CSV files
        elif ext == 'csv':
            try:
                csv_file = io.BytesIO(file_bytes)
                df = pd.read_csv(csv_file, encoding='utf-8', errors='ignore')
                extracted_text = df.to_string()
                if extracted_text.strip():
                    print(f"✅ Extracted {len(extracted_text)} characters from CSV: {filename}")
                    return extracted_text.strip()
            except Exception as e:
                print(f"⚠️ CSV extraction failed for {filename}: {e}")
        
        # JSON files
        elif ext == 'json':
            try:
                json_data = json.loads(file_bytes.decode('utf-8', errors='ignore'))
                extracted_text = json.dumps(json_data, indent=2, ensure_ascii=False)
                if extracted_text.strip():
                    print(f"✅ Extracted {len(extracted_text)} characters from JSON: {filename}")
                    return extracted_text.strip()
            except Exception as e:
                print(f"⚠️ JSON extraction failed for {filename}: {e}")
        
        # Plain text files (TXT, MD, etc.)
        elif ext in TEXT_SUMMARY_EXTENSIONS or (content_type and content_type.startswith('text/')):
            try:
                # Try UTF-8 first, then fallback to other encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        extracted_text = file_bytes.decode(encoding, errors='ignore')
                        if extracted_text.strip():
                            print(f"✅ Extracted {len(extracted_text)} characters from text file: {filename}")
                            return extracted_text.strip()
                    except:
                        continue
            except Exception as e:
                print(f"⚠️ Text extraction failed for {filename}: {e}")
        
        # Fallback: try to decode as text
        try:
            sample = file_bytes[:500_000]  # Limit to first ~500KB
            extracted_text = sample.decode('utf-8', errors='ignore')
            if extracted_text.strip() and len(extracted_text.strip()) > 50:
                print(f"✅ Extracted {len(extracted_text)} characters using fallback method: {filename}")
                return extracted_text.strip()
        except:
            pass
        
        print(f"⚠️ Could not extract text from {filename} (unsupported or binary file)")
        return ''
        
    except Exception as e:
        print(f"❌ Error extracting text from {filename}: {e}")
        import traceback
        traceback.print_exc()
        return ''


def extract_text_for_summary(file_bytes):
    """Legacy function - kept for backward compatibility"""
    if not file_bytes:
        return ''
    try:
        sample = file_bytes[:200_000]  # Limit to first ~200KB
        text = sample.decode('utf-8', errors='ignore')
        return text.strip()
    except Exception as exc:
        print(f"Failed to decode file content for summary: {exc}")
        return ''


AUTHOR_PATTERNS = [
    re.compile(r'\b(?:author|written by|prepared by|created by)\b[:\-]?\s*(?P<name>.+)', re.IGNORECASE),
    re.compile(r'^\s*by\s+(?P<name>.+)', re.IGNORECASE),
]

DATE_PATTERNS = [
    re.compile(r'\b\d{4}-\d{1,2}-\d{1,2}\b'),
    re.compile(r'\b\d{1,2}/\d{1,2}/\d{4}\b'),
    re.compile(r'\b\d{1,2}-\d{1,2}-\d{4}\b'),
    re.compile(r'\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|'
               r'jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+'
               r'\d{1,2},?\s+\d{4}\b', re.IGNORECASE),
    re.compile(r'\b\d{1,2}\s+(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|'
               r'jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+\d{4}\b',
               re.IGNORECASE),
    re.compile(r'\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|'
               r'jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+\d{4}\b',
               re.IGNORECASE),
]


def extract_metadata_from_text(text):
    """
    Attempt to extract publish date and author name from free-form document text.
    Returns a tuple of (publish_date, author_name) where either value may be None.
    """
    if not text:
        return None, None

    publish_date = None
    author_name = None

    # Limit processing to avoid scanning entire large documents repeatedly
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    limited_lines = lines[:200]  # Only inspect first 200 non-empty lines

    for line in limited_lines:
        if not author_name:
            for pattern in AUTHOR_PATTERNS:
                match = pattern.search(line)
                if match:
                    candidate = match.group('name').strip(" :-\t")
                    candidate = re.split(r'[|,;]|\s{2,}', candidate)[0].strip()
                    if candidate and candidate.lower() not in ('unknown', 'n/a'):
                        author_name = candidate
                        logger.debug("Extracted author from line '%s': %s", line, author_name)
                        break

        if not publish_date:
            for pattern in DATE_PATTERNS:
                match = pattern.search(line)
                if match:
                    candidate = match.group(0).strip(" :-\t")
                    if candidate:
                        publish_date = candidate
                        logger.debug("Extracted publish date from line '%s': %s", line, publish_date)
                        break

        if author_name and publish_date:
            break

    # If publish date still not found, search entire text for any date pattern
    if not publish_date:
        for pattern in DATE_PATTERNS:
            match = pattern.search(text)
            if match:
                publish_date = match.group(0).strip()
                logger.debug("Extracted publish date from document body: %s", publish_date)
                break

    return publish_date, author_name


def generate_file_summary(filename, file_bytes=None, content_type=None, extracted_text=None):
    """
    Generate a summary for uploaded files using OpenAI
    
    Args:
        filename: Name of the file
        file_bytes: File content as bytes (optional if extracted_text is provided)
        content_type: MIME type of the file (optional)
        extracted_text: Pre-extracted text from the file (optional, avoids re-extraction)
    """
    if not OPENAI_API_KEY:
        print(f"⚠️ OpenAI API key not set, skipping summary for {filename}")
        return None

    # Use provided extracted text or extract from file_bytes
    if extracted_text:
        text_content = extracted_text
        print(f"📄 Using pre-extracted text for {filename} ({len(text_content)} characters)")
    elif file_bytes:
        print(f"📄 Extracting text from {filename}...")
        text_content = extract_text_from_file(filename, file_bytes, content_type)
    else:
        print(f"⚠️ No text content provided for {filename}")
        return None

    if not text_content or len(text_content.strip()) < 10:
        print(f"⚠️ Could not extract sufficient text content from {filename} for summary")
        return None

    truncated_text = text_content[:MAX_SUMMARY_INPUT_CHARS]

    system_prompt = (
        "You create concise knowledge-base summaries for uploaded company documents. "
        "Write exactly three lines, each a short sentence under 120 characters, "
        "capturing the document's purpose, key points, and any notable actions."
    )

    user_prompt = (
        f"File name: {filename}\n\n"
        "Provide a three-line summary of the key ideas in this document snippet:\n"
        f"{truncated_text}"
    )

    try:
        print(f"🤖 Generating summary for {filename} using {OPENAI_SUMMARY_MODEL}...")
        if _openai_client:
            response = _openai_client.chat.completions.create(
                model=OPENAI_SUMMARY_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.2,
                max_tokens=220,
            )
            summary_text = response.choices[0].message.content.strip()
        else:
            if not openai:
                print(f"⚠️ OpenAI library not available for {filename}")
                return None
            response = openai.ChatCompletion.create(
                model=OPENAI_SUMMARY_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.2,
                max_tokens=220,
            )
            summary_text = response['choices'][0]['message']['content'].strip()

        if summary_text:
            print(f"✅ Successfully generated summary for {filename}")
            return summary_text
        else:
            print(f"⚠️ Empty summary response for {filename}")
    except Exception as exc:
        error_msg = str(exc)
        # Check for API key errors
        if '401' in error_msg or 'invalid_api_key' in error_msg or 'AuthenticationError' in str(type(exc)):
            print(f"❌ OpenAI API key error for {filename}: Please check your OPENAI_API_KEY environment variable")
            print(f"   Error: Invalid or missing API key. Update your .env file with a valid OpenAI API key.")
        else:
            print(f"❌ Summary generation failed for {filename}: {exc}")
            import traceback
            traceback.print_exc()

    return None


def analyze_document_with_assistants_api(filename, file_bytes, content_type=None):
    """
    Analyze document using OpenAI Assistants API with File Search.
    Extracts: summary, publish date, and author name.
    
    Args:
        filename: Name of the file
        file_bytes: File content as bytes
        content_type: MIME type of the file (optional)
    
    Returns:
        dict with keys: summary, publish_date, author_name
        Returns None if analysis fails
    """
    if not OPENAI_API_KEY or not _openai_client:
        print(f"⚠️ OpenAI API key or client not available, skipping Assistants API analysis for {filename}")
        return None
    
    try:
        print(f"🤖 Starting Assistants API analysis for {filename}...")
        
        # Step 1: Upload file to OpenAI
        print(f"📤 Step 1: Uploading {filename} to OpenAI...")
        file_obj = _openai_client.files.create(
            file=(filename, io.BytesIO(file_bytes), content_type or 'application/octet-stream'),
            purpose="assistants"
        )
        file_id = file_obj.id
        print(f"✅ File uploaded: {file_id}")
        
        # Step 2: Create an assistant with file search capability
        print(f"🔧 Step 2: Creating assistant...")
        assistant = _openai_client.beta.assistants.create(
            name="Document Analyzer",
            instructions=(
                "You are a helpful assistant that analyzes documents and extracts key information. "
                "When asked about a document, provide: "
                "1. A concise summary (3 lines, each under 120 characters) "
                "2. The publish date (if found, otherwise 'unknown') "
                "3. The author name (if found, otherwise 'unknown')"
            ),
            model="gpt-4o-mini",
            tools=[{"type": "file_search"}]
        )
        assistant_id = assistant.id
        print(f"✅ Assistant created: {assistant_id}")
        
        # Step 3: Create a thread with the document attached
        print(f"💬 Step 3: Creating thread with document...")
        thread = _openai_client.beta.threads.create(
            messages=[
                {
                    "role": "user",
                    "content": (
                        f"Please analyze this document ({filename}) and provide the following information:\n"
                        "1. A concise summary (exactly 3 lines, each under 120 characters) covering the document's purpose, key points, and notable actions.\n"
                        "2. The publish date (if found in the document, format as YYYY-MM-DD, otherwise respond with 'unknown').\n"
                        "3. The author name (if found in the document, otherwise respond with 'unknown').\n\n"
                        "Please format your response as:\n"
                        "SUMMARY:\n[three-line summary]\n\n"
                        "PUBLISH_DATE:\n[date or unknown]\n\n"
                        "AUTHOR:\n[name or unknown]"
                    ),
                    "attachments": [
                        {
                            "file_id": file_id,
                            "tools": [{"type": "file_search"}]
                        }
                    ]
                }
            ]
        )
        thread_id = thread.id
        print(f"✅ Thread created: {thread_id}")
        
        # Step 4: Run the assistant
        print(f"🔄 Step 4: Running assistant...")
        run = _openai_client.beta.threads.runs.create_and_poll(
            thread_id=thread_id,
            assistant_id=assistant_id,
            timeout=300  # 5 minutes timeout
        )
        
        if run.status != 'completed':
            print(f"⚠️ Run status: {run.status}")
            if run.status == 'failed':
                print(f"❌ Run failed: {run.last_error}")
            # Clean up
            try:
                _openai_client.files.delete(file_id)
                _openai_client.beta.assistants.delete(assistant_id)
            except:
                pass
            return None
        
        # Step 5: Get the response
        print(f"📥 Step 5: Retrieving response...")
        messages = _openai_client.beta.threads.messages.list(thread_id=thread_id)
        
        if not messages.data or len(messages.data) == 0:
            print(f"⚠️ No messages in thread")
            # Clean up
            try:
                _openai_client.files.delete(file_id)
                _openai_client.beta.assistants.delete(assistant_id)
            except:
                pass
            return None
        
        # Get the assistant's response (first message should be from assistant)
        assistant_message = None
        for msg in messages.data:
            if msg.role == 'assistant' and msg.content:
                assistant_message = msg
                break
        
        if not assistant_message or not assistant_message.content:
            print(f"⚠️ No assistant message found")
            # Clean up
            try:
                _openai_client.files.delete(file_id)
                _openai_client.beta.assistants.delete(assistant_id)
            except:
                pass
            return None
        
        # Extract text from response
        response_text = ""
        for content_block in assistant_message.content:
            if hasattr(content_block, 'text') and hasattr(content_block.text, 'value'):
                response_text = content_block.text.value
                break
            elif isinstance(content_block, dict) and 'text' in content_block:
                response_text = content_block['text'].get('value', '')
                break
        
        if not response_text:
            print(f"⚠️ Empty response text")
            # Clean up
            try:
                _openai_client.files.delete(file_id)
                _openai_client.beta.assistants.delete(assistant_id)
            except:
                pass
            return None
        
        print(f"✅ Received response: {response_text[:200]}...")
        
        # Step 6: Parse the response
        summary = None
        publish_date = "unknown"
        author_name = "unknown"
        
        # Parse SUMMARY section
        if "SUMMARY:" in response_text:
            summary_section = response_text.split("SUMMARY:")[1]
            if "PUBLISH_DATE:" in summary_section:
                summary_section = summary_section.split("PUBLISH_DATE:")[0]
            summary_lines = [line.strip() for line in summary_section.strip().split("\n") if line.strip()]
            if summary_lines:
                summary = "\n".join(summary_lines[:3])  # Take first 3 lines
        
        # Parse PUBLISH_DATE section
        if "PUBLISH_DATE:" in response_text:
            date_section = response_text.split("PUBLISH_DATE:")[1]
            if "AUTHOR:" in date_section:
                date_section = date_section.split("AUTHOR:")[0]
            date_value = date_section.strip().split("\n")[0].strip()
            if date_value and date_value.lower() != "unknown":
                publish_date = date_value
            else:
                publish_date = "unknown"
        
        # Parse AUTHOR section
        if "AUTHOR:" in response_text:
            author_section = response_text.split("AUTHOR:")[1]
            author_value = author_section.strip().split("\n")[0].strip()
            if author_value and author_value.lower() != "unknown":
                author_name = author_value
            else:
                author_name = "unknown"
        
        # If summary wasn't found in structured format, try to extract it from the response
        if not summary:
            # Try to get first 3 lines as summary
            lines = [line.strip() for line in response_text.strip().split("\n") if line.strip()]
            if lines:
                summary = "\n".join(lines[:3])
        
        # Clean up: Delete file and assistant
        try:
            print(f"🧹 Cleaning up OpenAI resources...")
            _openai_client.files.delete(file_id)
            _openai_client.beta.assistants.delete(assistant_id)
            print(f"✅ Cleanup completed")
        except Exception as cleanup_error:
            print(f"⚠️ Cleanup warning: {cleanup_error}")
        
        result = {
            'summary': summary or SUMMARY_PLACEHOLDER,
            'publish_date': publish_date,
            'author_name': author_name
        }
        
        print(f"✅ Analysis complete for {filename}:")
        print(f"   Summary: {summary[:100] if summary else 'N/A'}...")
        print(f"   Publish Date: {publish_date}")
        print(f"   Author: {author_name}")
        
        return result
        
    except Exception as exc:
        error_msg = str(exc)
        print(f"❌ Assistants API analysis failed for {filename}: {exc}")
        import traceback
        traceback.print_exc()
        
        # Try to clean up on error
        try:
            if 'file_id' in locals():
                _openai_client.files.delete(file_id)
            if 'assistant_id' in locals():
                _openai_client.beta.assistants.delete(assistant_id)
        except:
            pass
        
        return None


def save_file_metadata(
    blob_path,
    filename,
    author_email,
    summary_text,
    file_size=None,
    content_type=None,
    author_name=None,
    publish_date=None,
    document_author=None
):
    if not blob_path:
        return

    try:
        collection = get_mongo_connection('file-metadata')
    except Exception as exc:
        print(f"Unable to connect to MongoDB for metadata storage: {exc}")
        return

    now = datetime.now(pytz.UTC)
    summary_to_store = (summary_text or SUMMARY_PLACEHOLDER).strip()
    preferred_author_name = author_name or author_email
    # Use document_author from analysis if available, otherwise use uploader's name
    final_document_author = document_author if document_author and document_author.lower() != 'unknown' else None
    final_publish_date = publish_date if publish_date and publish_date.lower() != 'unknown' else None

    try:
        update_data = {
            'file_name': filename,
            'author_email': author_email,
            'author_name': preferred_author_name,
            'summary': summary_to_store,
            'size': file_size,
            'content_type': content_type,
            'updated_at': now,
            'summary_model': OPENAI_SUMMARY_MODEL if summary_text else None
        }
        
        # Add document metadata if available
        if final_document_author:
            update_data['document_author'] = final_document_author
        if final_publish_date:
            update_data['publish_date'] = final_publish_date
        
        collection.update_one(
            {'blob_path': blob_path},
            {
                '$set': update_data,
                '$setOnInsert': {
                    'blob_path': blob_path,
                    'uploaded_at': now
                }
            },
            upsert=True
        )
        print(f"Stored metadata for {blob_path}")
        if final_document_author:
            print(f"   Document author: {final_document_author}")
        if final_publish_date:
            print(f"   Publish date: {final_publish_date}")
    except Exception as exc:
        print(f"Failed to store metadata for {filename}: {exc}")


def fetch_file_metadata_map(file_paths):
    if not file_paths:
        return {}

    try:
        collection = get_mongo_connection('file-metadata')
    except Exception as exc:
        print(f"Unable to connect to MongoDB for metadata retrieval: {exc}")
        return {}

    try:
        # First, try to match by blob_path (exact match)
        docs_by_path = collection.find({'blob_path': {'$in': file_paths}})
        metadata_map = {}
        matched_paths = set()
        
        # Process exact blob_path matches
        for doc in docs_by_path:
            blob_path = doc.get('blob_path')
            if not blob_path:
                continue
            
            matched_paths.add(blob_path)
            
            uploaded_at = doc.get('uploaded_at')
            uploaded_at_iso = None
            if isinstance(uploaded_at, datetime):
                if uploaded_at.tzinfo is None:
                    uploaded_at = uploaded_at.replace(tzinfo=pytz.UTC)
                uploaded_at_iso = uploaded_at.isoformat()
            elif uploaded_at:
                uploaded_at_iso = str(uploaded_at)

            metadata_map[blob_path] = {
                'summary': doc.get('summary') or SUMMARY_PLACEHOLDER,
                'author_name': doc.get('author_name') or doc.get('author_email'),
                'author_email': doc.get('author_email'),
                'uploaded_at': uploaded_at_iso,
                'file_name': doc.get('file_name'),
                'size': doc.get('size'),
                'content_type': doc.get('content_type'),
                'publish_date': doc.get('publish_date'),
                'document_author': doc.get('document_author'),
            }
            print(f"✅ Matched by blob_path: {blob_path}")
        
        # For files not matched by blob_path, try matching by file_name
        unmatched_paths = [path for path in file_paths if path not in matched_paths]
        if unmatched_paths:
            # Extract filenames from unmatched paths
            file_names = [path.split('/')[-1] if '/' in path else path for path in unmatched_paths]
            
            # Query by file_name
            docs_by_name = collection.find({'file_name': {'$in': file_names}})
            
            # Create a mapping of filename to full paths
            filename_to_paths = {}
            for path in unmatched_paths:
                filename = path.split('/')[-1] if '/' in path else path
                if filename not in filename_to_paths:
                    filename_to_paths[filename] = []
                filename_to_paths[filename].append(path)
            
            # Process file_name matches
            for doc in docs_by_name:
                file_name = doc.get('file_name')
                if not file_name:
                    continue
                
                # Match this document to all paths with this filename
                matching_paths = filename_to_paths.get(file_name, [])
                for path in matching_paths:
                    if path not in metadata_map:  # Only if not already matched
                        uploaded_at = doc.get('uploaded_at')
                        uploaded_at_iso = None
                        if isinstance(uploaded_at, datetime):
                            if uploaded_at.tzinfo is None:
                                uploaded_at = uploaded_at.replace(tzinfo=pytz.UTC)
                            uploaded_at_iso = uploaded_at.isoformat()
                        elif uploaded_at:
                            uploaded_at_iso = str(uploaded_at)
                        
                        metadata_map[path] = {
                            'summary': doc.get('summary') or SUMMARY_PLACEHOLDER,
                            'author_name': doc.get('author_name') or doc.get('author_email'),
                            'author_email': doc.get('author_email'),
                            'uploaded_at': uploaded_at_iso,
                            'file_name': doc.get('file_name'),
                            'size': doc.get('size'),
                            'content_type': doc.get('content_type'),
                            'publish_date': doc.get('publish_date'),
                            'document_author': doc.get('document_author'),
                        }
                        print(f"✅ Matched by file_name: {file_name} -> {path}")
        
        # Log unmatched files
        all_matched = set(metadata_map.keys())
        unmatched = [path for path in file_paths if path not in all_matched]
        if unmatched:
            print(f"⚠️ No metadata found for {len(unmatched)} file(s): {unmatched[:5]}...")  # Show first 5
        
        print(f"📊 Metadata matching summary: {len(metadata_map)}/{len(file_paths)} files matched")
        return metadata_map
    except Exception as exc:
        print(f"Failed to fetch file metadata: {exc}")
        import traceback
        traceback.print_exc()
        return {}





# testing 
# testing 
@app.route('/upload/multiple', methods=['POST'])
def handle_upload():
    """Handle file upload and forward to Azure service"""
    try:
        if 'email' not in session:
            return jsonify({
                'status': 'error',
                'message': 'Unauthorized'
            }), 401

        author_email = session.get('email')
        author_name = (
            session.get('display_name')
            or session.get('name')
            or session.get('full_name')
            or author_email
        )

        # Check if files are in the request
        if 'files' not in request.files:
            return jsonify({
                'status': 'error',
                'message': 'No files selected'
            }), 400

        files = request.files.getlist('files')
        
        if not files or files[0].filename == '':
            return jsonify({
                'status': 'error',
                'message': 'No files selected'
            }), 400

        uploaded_files = []
        errors = []
        total_chunks = 0

        # Azure service base endpoint
        azure_base_url = f"{BASE_URL}/upload/multipdf-file"

        for file in files:
            if file and allowed_file(file.filename):
                # Check file size
                file.seek(0, os.SEEK_END)
                file_size = file.tell()
                file.seek(0)  # Reset file pointer
                
                if file_size > MAX_FILE_SIZE:
                    errors.append(f'{file.filename}: File size exceeds 10MB limit')
                    uploaded_files.append({
                        'filename': file.filename,
                        'error': 'File size exceeds 10MB limit'
                    })
                    continue

                # Secure the filename
                filename = secure_filename(file.filename)
                extracted_text = None
                
                try:
                    # ========== STEP 1: READ FILE CONTENT ==========
                    file.seek(0)  # Reset pointer before reading
                    file_content = file.read()
                    
                    # Determine MIME type based on filename; default to octet-stream
                    guessed_mime, _ = mimetypes.guess_type(filename)
                    content_type = guessed_mime or 'application/octet-stream'
                    
                    print(f"📄 Processing {filename}...")
                    print(f"   File size: {file_size} bytes")
                    print(f"   Content type: {content_type}")
                    logger.info("Processing upload '%s' (size=%d, content_type=%s)", filename, file_size, content_type)
                    
                    # ========== STEP 2: ANALYZE DOCUMENT USING ASSISTANTS API ==========
                    print(f"🤖 Step 1: Analyzing {filename} using OpenAI Assistants API...")
                    
                    # Initialize variables
                    summary_text = None
                    publish_date = None
                    document_author = None
                    
                    # Use Assistants API to extract summary, publish date, and author
                    analysis_result = analyze_document_with_assistants_api(
                        filename=filename,
                        file_bytes=file_content,
                        content_type=content_type
                    )
                    
                    # Extract results from analysis
                    if analysis_result:
                        summary_text = analysis_result.get('summary')
                        publish_date = analysis_result.get('publish_date')
                        document_author = analysis_result.get('author_name')
                        
                        print(f"✅ Analysis complete for {filename}:")
                        print(f"   Summary: {'Generated' if summary_text and summary_text != SUMMARY_PLACEHOLDER else 'Placeholder'}")
                        print(f"   Publish Date: {publish_date}")
                        print(f"   Document Author: {document_author}")
                        logger.info(
                            "Assistants API metadata for '%s' -> publish_date=%s, author=%s",
                            filename,
                            publish_date,
                            document_author,
                        )
                    else:
                        # Fallback to old method if Assistants API fails
                        print(f"⚠️ Assistants API analysis failed, falling back to text extraction...")
                        extracted_text = extract_text_from_file(filename, file_content, content_type)
                        
                        if not extracted_text or len(extracted_text.strip()) < 10:
                            print(f"⚠️ Could not extract text from {filename}, proceeding without summary")
                            summary_text = None
                        else:
                            print(f"✅ Extracted {len(extracted_text)} characters from {filename}")
                            logger.info(
                                "Extracted %d characters from '%s' for offline analysis",
                                len(extracted_text),
                                filename,
                            )
                            
                            # Generate summary using old method
                            print(f"🤖 Generating summary for {filename} using OpenAI...")
                            summary_text = generate_file_summary(
                                filename, 
                                file_bytes=file_content, 
                                content_type=content_type,
                                extracted_text=extracted_text
                            )
                            
                            if summary_text:
                                print(f"✅ Summary generated: {summary_text[:100]}...")
                                logger.info("Generated fallback summary for '%s'", filename)
                            else:
                                print(f"⚠️ No summary generated for {filename}, using placeholder")

                    # If publish date or document author are missing, try extracting from text content
                    needs_publish_date = not publish_date or str(publish_date).strip().lower() == 'unknown'
                    needs_author = not document_author or str(document_author).strip().lower() == 'unknown'

                    if needs_publish_date or needs_author:
                        if extracted_text is None:
                            extracted_text = extract_text_from_file(filename, file_content, content_type)

                        if extracted_text:
                            derived_publish_date, derived_author = extract_metadata_from_text(extracted_text)

                            if needs_publish_date and derived_publish_date:
                                publish_date = derived_publish_date

                            if needs_author and derived_author:
                                document_author = derived_author

                            if derived_publish_date or derived_author:
                                logger.info(
                                    "Resolved metadata for '%s' via text parsing -> publish_date=%s, author=%s",
                                    filename,
                                    publish_date,
                                    document_author,
                                )
                        else:
                            logger.debug(
                                "No extracted text available to derive metadata for '%s'",
                                filename,
                            )
                    
                    # ========== STEP 4: UPLOAD TO BLOB STORAGE ==========
                    print(f"📤 Step 3: Uploading {filename} to Azure blob storage...")
                    print(f"   URL: {azure_base_url}")
                    
                    # IMPORTANT: Azure expects 'file' (singular) as the parameter name
                    files_to_upload = {
                        'file': (filename, file_content, content_type)
                    }
                    
                    # Upload to Azure service
                    response = requests.post(
                        azure_base_url,
                        files=files_to_upload,
                        timeout=120
                    )
                    
                    print(f"📥 Azure response status: {response.status_code}")
                    print(f"📥 Azure response: {response.text}")
                    
                    # Try to parse JSON response
                    try:
                        azure_response = response.json()
                        print(f"📥 Azure response JSON: {azure_response}")
                    except ValueError as json_error:
                        # If response is not JSON, treat it as an error
                        print(f"❌ Azure returned non-JSON response: {response.text}")
                        errors.append(f'{filename}: Invalid response from server')
                        uploaded_files.append({
                            'filename': filename,
                            'error': f'Invalid server response (Status: {response.status_code})'
                        })
                        continue
                    
                    # Check if upload was successful
                    if response.status_code == 200:
                        # Check if Azure response indicates success
                        if azure_response.get('status') == 'ok':
                            # Extract file path and calculate chunks (you may need to adjust this)
                            file_path = azure_response.get('path', '')
                            message = azure_response.get('message', '')
                            
                            # Estimate chunks based on file size (adjust as needed)
                            # Assuming each chunk is ~1000 characters or 1KB
                            chunks = file_size // 1024 if file_size > 0 else 1
                            total_chunks += chunks
                            
                            # ========== STEP 5: SAVE METADATA WITH SUMMARY ==========
                            print(f"💾 Step 4: Saving metadata for {filename}...")
                            save_file_metadata(
                                blob_path=file_path,
                                filename=azure_response.get('filename', filename),
                                author_email=author_email,
                                summary_text=summary_text,
                                file_size=file_size,
                                content_type=content_type,
                                author_name=author_name,
                                publish_date=publish_date,
                                document_author=document_author
                            )
                            print(f"✅ Metadata saved for {filename}")

                            uploaded_files.append({
                                'filename': azure_response.get('filename', filename),
                                'size': file_size,
                                'chunks': chunks,
                                'blob_path': file_path,
                                'summary': summary_text or SUMMARY_PLACEHOLDER,
                                'author': author_name,
                                'uploaded_at': datetime.now(pytz.UTC).isoformat(),
                                'publishDate': publish_date,
                                'publish_date': publish_date,
                                'documentAuthor': document_author,
                                'document_author': document_author,
                            })
                            print(f"✅ Successfully processed {filename}")
                            print(f"   Path: {file_path}")
                            print(f"   Message: {message}")
                            print(f"   Summary: {'Generated' if summary_text else 'Placeholder'}")
                            print(f"   Next: File will be indexed for embeddings in ingestion step")
                        else:
                            # Azure returned 200 but with error status
                            error_msg = azure_response.get('message', azure_response.get('error', 'Unknown error'))
                            errors.append(f'{filename}: {error_msg}')
                            uploaded_files.append({
                                'filename': filename,
                                'error': error_msg
                            })
                            print(f"⚠️ Azure returned error for {filename}: {error_msg}")
                    else:
                        # Non-200 status code
                        error_msg = azure_response.get('message', azure_response.get('error', f'HTTP {response.status_code}'))
                        errors.append(f'{filename}: {error_msg}')
                        uploaded_files.append({
                            'filename': filename,
                            'error': f'{error_msg} (Status: {response.status_code})'
                        })
                        print(f"❌ Azure upload failed for {filename}: {error_msg}")
                        
                except requests.exceptions.Timeout:
                    error_msg = "Upload timeout - file may be too large or server is slow"
                    print(f"⏱️ Timeout error for {filename}")
                    errors.append(f'{filename}: {error_msg}')
                    uploaded_files.append({
                        'filename': filename,
                        'error': error_msg
                    })
                except requests.exceptions.ConnectionError as conn_error:
                    error_msg = f"Connection error - cannot reach Azure server"
                    print(f"🔌 Connection error for {filename}: {str(conn_error)}")
                    errors.append(f'{filename}: {error_msg}')
                    uploaded_files.append({
                        'filename': filename,
                        'error': error_msg
                    })
                except requests.exceptions.RequestException as req_error:
                    error_msg = f"Network error - {str(req_error)}"
                    print(f"🌐 Network error for {filename}: {str(req_error)}")
                    errors.append(f'{filename}: {error_msg}')
                    uploaded_files.append({
                        'filename': filename,
                        'error': error_msg
                    })
                except Exception as upload_error:
                    error_msg = f"Upload error - {str(upload_error)}"
                    print(f"❌ Unexpected error for {filename}: {str(upload_error)}")
                    import traceback
                    traceback.print_exc()
                    errors.append(f'{filename}: {error_msg}')
                    uploaded_files.append({
                        'filename': filename,
                        'error': error_msg
                    })
            else:
                errors.append(f"{file.filename}: Invalid file type")
                uploaded_files.append({
                    'filename': file.filename,
                    'error': 'Invalid file type'
                })

        # Determine overall status
        successful_uploads = [f for f in uploaded_files if 'error' not in f]
        failed_uploads = [f for f in uploaded_files if 'error' in f]
        
        print(f"\n📊 Upload Summary:")
        print(f"   Total files: {len(uploaded_files)}")
        print(f"   Successful: {len(successful_uploads)}")
        print(f"   Failed: {len(failed_uploads)}")
        print(f"   Total chunks: {total_chunks}")

        # Return response matching the frontend expectations
        if len(successful_uploads) > 0:
            return jsonify({
                'status': 'ok',
                'message': f'Successfully uploaded {len(successful_uploads)} of {len(uploaded_files)} file(s)',
                'total_files': len(uploaded_files),
                'successful_uploads': len(successful_uploads),
                'failed_uploads': len(failed_uploads),
                'total_chunks': total_chunks,
                'uploaded_files': uploaded_files,
                'index': 'gptindex',
                'folder_path': 'content',
                'errors': errors if errors else None
            }), 200
        else:
            # All uploads failed
            return jsonify({
                'status': 'error',
                'message': 'All file uploads failed',
                'total_files': len(uploaded_files),
                'uploaded_files': uploaded_files,
                'errors': errors
            }), 400

    except Exception as e:
        print(f"❌ Critical error in upload handler: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'status': 'error',
            'message': f'Server error: {str(e)}'
        }), 500
# ingestion
@app.route('/ingest/files', methods=['POST'])
def ingest_files():
    """Handle file ingestion after upload"""
    try:
        if 'email' not in session:
            return jsonify({
                'status': 'error',
                'message': 'Unauthorized'
            }), 401

        # Get data from request body
        data = request.get_json()
        
        if not data:
            return jsonify({
                'status': 'error',
                'message': 'No data provided'
            }), 400
        
        # Extract file paths
        file_paths = data.get('file_paths', [])
        
        if not file_paths or len(file_paths) == 0:
            return jsonify({
                'status': 'error',
                'message': 'No file paths provided'
            }), 400
        
        print(f"🔄 Starting ingestion for {len(file_paths)} file(s)")
        print(f"   File paths: {file_paths}")
        
        # Azure ingestion endpoint
        ingestion_url = f"{BASE_URL}/ingest/multipdf-file"
        
        ingestion_results = []
        errors = []
        
        # Send each file path for ingestion
        for file_path in file_paths:
            try:
                print(f"📥 Ingesting file: {file_path}")
                
                # Prepare the request body
                ingestion_payload = {
                    "path": file_path
                }
                
                # Send to Azure ingestion service
                response = requests.post(
                    ingestion_url,
                    json=ingestion_payload,
                    headers={'Content-Type': 'application/json'},
                    timeout=180  # 3 minutes timeout for ingestion
                )
                
                print(f"📥 Ingestion response status: {response.status_code}")
                print(f"📥 Ingestion response: {response.text}")
                
                # Parse response
                try:
                    ingestion_response = response.json()
                except ValueError:
                    print(f"❌ Non-JSON response from ingestion service")
                    errors.append(f'{file_path}: Invalid response from ingestion service')
                    ingestion_results.append({
                        'file_path': file_path,
                        'status': 'error',
                        'error': 'Invalid response from ingestion service'
                    })
                    continue
                
                # Check if ingestion was successful
                if response.status_code == 200:
                    if ingestion_response.get('status') == 'ok' or ingestion_response.get('success') == True:
                        ingestion_results.append({
                            'file_path': file_path,
                            'status': 'success',
                            'message': ingestion_response.get('message', 'File ingested successfully'),
                            'data': ingestion_response
                        })
                        print(f"✅ Successfully ingested: {file_path}")
                    else:
                        error_msg = ingestion_response.get('message', ingestion_response.get('error', 'Unknown error'))
                        errors.append(f'{file_path}: {error_msg}')
                        ingestion_results.append({
                            'file_path': file_path,
                            'status': 'error',
                            'error': error_msg
                        })
                        print(f"⚠️ Ingestion error for {file_path}: {error_msg}")
                else:
                    error_msg = ingestion_response.get('message', ingestion_response.get('error', f'HTTP {response.status_code}'))
                    errors.append(f'{file_path}: {error_msg}')
                    ingestion_results.append({
                        'file_path': file_path,
                        'status': 'error',
                        'error': error_msg
                    })
                    print(f"❌ Ingestion failed for {file_path}: {error_msg}")
                    
            except requests.exceptions.Timeout:
                error_msg = "Ingestion timeout"
                print(f"⏱️ Timeout error for {file_path}")
                errors.append(f'{file_path}: {error_msg}')
                ingestion_results.append({
                    'file_path': file_path,
                    'status': 'error',
                    'error': error_msg
                })
            except requests.exceptions.RequestException as req_error:
                error_msg = f"Network error - {str(req_error)}"
                print(f"🌐 Network error for {file_path}: {str(req_error)}")
                errors.append(f'{file_path}: {error_msg}')
                ingestion_results.append({
                    'file_path': file_path,
                    'status': 'error',
                    'error': error_msg
                })
            except Exception as e:
                error_msg = f"Ingestion error - {str(e)}"
                print(f"❌ Error for {file_path}: {str(e)}")
                errors.append(f'{file_path}: {error_msg}')
                ingestion_results.append({
                    'file_path': file_path,
                    'status': 'error',
                    'error': error_msg
                })
        
        # Count successful and failed ingestions
        successful_ingestions = [r for r in ingestion_results if r.get('status') == 'success']
        failed_ingestions = [r for r in ingestion_results if r.get('status') == 'error']
        
        print(f"\n📊 Ingestion Summary:")
        print(f"   Total files: {len(ingestion_results)}")
        print(f"   Successful: {len(successful_ingestions)}")
        print(f"   Failed: {len(failed_ingestions)}")
        
        # Return response
        if len(successful_ingestions) > 0:
            return jsonify({
                'status': 'ok',
                'message': f'Successfully ingested {len(successful_ingestions)} of {len(ingestion_results)} file(s)',
                'total_files': len(ingestion_results),
                'successful_ingestions': len(successful_ingestions),
                'failed_ingestions': len(failed_ingestions),
                'ingestion_results': ingestion_results,
                'errors': errors if errors else None
            }), 200
        else:
            return jsonify({
                'status': 'error',
                'message': 'All file ingestions failed',
                'total_files': len(ingestion_results),
                'ingestion_results': ingestion_results,
                'errors': errors
            }), 400
            
    except Exception as e:
        print(f"❌ Critical error in ingestion handler: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'status': 'error',
            'message': f'Server error: {str(e)}'
        }), 500



@app.route('/api/get-files', methods=['GET'])
def get_files():
    """
    Proxy endpoint to fetch files from backend service
    Returns files grouped by email/ID, including ungrouped files under "Ungrouped" category
    Now includes file metadata (size, dates, content type)
    
    Query Parameters:
    - sort_by: 'name', 'size', 'date', 'type' (default: 'date')
    - sort_order: 'asc', 'desc' (default: 'desc')
    - email_filter: filter by specific email/ID
    """
    try:
        # Get query parameters
        sort_by = request.args.get('sort_by', 'date')
        sort_order = request.args.get('sort_order', 'desc')
        email_filter = request.args.get('email_filter', None)
        
        # Backend service URL
        backend_url = f'{BASE_URL}/list_content_files'
        
        # Call the backend service with increased timeout
        try:
            response = requests.get(backend_url, timeout=30)
        except requests.exceptions.Timeout:
            return jsonify({
                'status': 'error',
                'message': 'Backend service request timed out. Please try again later.'
            }), 504
        except requests.exceptions.ConnectionError:
            return jsonify({
                'status': 'error',
                'message': 'Could not connect to backend service. Please check your connection.'
            }), 503
        except requests.exceptions.RequestException as e:
            return jsonify({
                'status': 'error',
                'message': f'Error connecting to backend service: {str(e)}'
            }), 500
        
        # Check if request was successful
        if response.status_code != 200:
            return jsonify({
                'status': 'error',
                'message': f'Backend service returned status code: {response.status_code}'
            }), response.status_code
        
        # Get the data from backend
        backend_data = response.json()
        
        # Process and group files by email
        files_by_email = {}
        ungrouped_files = []  # Files without email/ID prefix
        total_size = 0  # Track total size of all files
        all_file_entries = []
        
        if 'files' in backend_data and isinstance(backend_data['files'], list):
            for file_item in backend_data['files']:
                # Handle both old format (strings) and new format (objects with metadata)
                if isinstance(file_item, str):
                    # Old format: just file path as string
                    file_path = file_item
                    file_metadata = {
                        'fileName': file_path.split('/')[-1],
                        'fullPath': file_path,
                        'size': None,
                        'sizeFormatted': 'Unknown',
                        'createdOn': None,
                        'lastModified': None,
                        'contentType': 'Unknown',
                        'etag': None
                    }
                else:
                    # New format: object with metadata
                    file_path = file_item.get('name', '')
                    file_size = file_item.get('size', 0)
                    
                    file_metadata = {
                        'fileName': file_path.split('/')[-1],
                        'fullPath': file_path,
                        'size': file_size,
                        'sizeFormatted': format_file_size(file_size),
                        'createdOn': file_item.get('created_on'),
                        'lastModified': file_item.get('last_modified'),
                        'contentType': file_item.get('content_type', 'Unknown'),
                        'etag': file_item.get('etag')
                    }
                    
                    # Add to total size
                    if file_size:
                        total_size += file_size
                
                # Check if file path contains "/"
                if '/' in file_path:
                    parts = file_path.split('/', 1)  # Split only on first "/"
                    email = parts[0]  # e.g., "awais@gmail.com"
                    file_name = parts[1]  # e.g., "DLIMS.pdf"
                    
                    # Skip if email filter is specified and doesn't match
                    if email_filter and email != email_filter:
                        continue
                    
                    if email not in files_by_email:
                        files_by_email[email] = {
                            'files': [],
                            'totalSize': 0,
                            'fileCount': 0
                        }
                    
                    # Update fileName to be just the file name without email prefix
                    file_metadata['fileName'] = file_name
                    files_by_email[email]['files'].append(file_metadata)
                    all_file_entries.append(file_metadata)
                    files_by_email[email]['fileCount'] += 1
                    if file_metadata['size']:
                        files_by_email[email]['totalSize'] += file_metadata['size']
                else:
                    # File without email/ID prefix
                    if not email_filter or email_filter == 'Ungrouped':
                        ungrouped_files.append(file_metadata)
                        all_file_entries.append(file_metadata)
        
        # Sort files within each email group
        for email in files_by_email:
            files_by_email[email]['files'] = sort_files(
                files_by_email[email]['files'], 
                sort_by, 
                sort_order
            )
            # Add formatted total size
            files_by_email[email]['totalSizeFormatted'] = format_file_size(
                files_by_email[email]['totalSize']
            )
        
        # Sort ungrouped files
        ungrouped_files = sort_files(ungrouped_files, sort_by, sort_order)
        
        # Add ungrouped files to files_by_email under "Ungrouped" category
        if ungrouped_files:
            ungrouped_total_size = sum(f.get('size', 0) for f in ungrouped_files if f.get('size'))
            files_by_email['Ungrouped'] = {
                'files': ungrouped_files,
                'totalSize': ungrouped_total_size,
                'totalSizeFormatted': format_file_size(ungrouped_total_size),
                'fileCount': len(ungrouped_files)
            }

        # Attach metadata (summary, author, upload date)
        file_paths = [entry.get('fullPath') for entry in all_file_entries if entry.get('fullPath')]
        
        print(f"\n📋 Fetching metadata for {len(file_paths)} files from MongoDB collection 'file-metadata'")
        if file_paths:
            print(f"   Sample paths: {file_paths[:3]}...")  # Show first 3 paths
        
        metadata_map = fetch_file_metadata_map(file_paths)
        
        print(f"📋 Metadata map contains {len(metadata_map)} entries")
        if metadata_map:
            print(f"   Sample matched paths: {list(metadata_map.keys())[:3]}...")
        
        matched_count = 0
        unmatched_count = 0
        
        for entry in all_file_entries:
            full_path = entry.get('fullPath')
            file_name = entry.get('fileName', 'unknown')
            meta = metadata_map.get(full_path) if full_path else None
            
            if meta:
                matched_count += 1
                # Set summary (avoid placeholder if we have actual summary)
                summary_value = meta.get('summary')
                if summary_value and summary_value.strip() and summary_value.strip() != SUMMARY_PLACEHOLDER:
                    entry['summary'] = summary_value.strip()
                else:
                    entry['summary'] = SUMMARY_PLACEHOLDER
                
                # Set author fields
                entry['uploadedBy'] = meta.get('author_name') or meta.get('author_email') or 'Unknown author'
                entry['authorEmail'] = meta.get('author_email')
                entry['author_name'] = meta.get('author_name') or meta.get('author_email')
                entry['uploadedAt'] = meta.get('uploaded_at')
                
                # Set document metadata (publish date and document author)
                entry['publishDate'] = meta.get('publish_date')
                entry['documentAuthor'] = meta.get('document_author')
                
                # Debug log for files with summaries
                if entry['summary'] and entry['summary'] != SUMMARY_PLACEHOLDER:
                    print(f"✅ [{matched_count}] File: {file_name}")
                    print(f"   Path: {full_path}")
                    print(f"   Summary: {entry['summary'][:80]}...")
                    print(f"   Author: {entry['uploadedBy']}")
                    print(f"   Uploaded: {entry.get('uploadedAt', 'N/A')}")
                    if entry.get('publishDate'):
                        print(f"   Publish Date: {entry['publishDate']}")
                    if entry.get('documentAuthor'):
                        print(f"   Document Author: {entry['documentAuthor']}")
            else:
                unmatched_count += 1
                entry.setdefault('summary', SUMMARY_PLACEHOLDER)
                entry.setdefault('uploadedBy', 'Unknown author')
                entry.setdefault('publishDate', None)
                entry.setdefault('documentAuthor', None)
                if full_path:
                    print(f"⚠️ [{unmatched_count}] No metadata found for: {file_name} (path: {full_path})")
        
        print(f"\n📊 Final metadata attachment: {matched_count} matched, {unmatched_count} unmatched out of {len(all_file_entries)} total files\n")
        
        # Return processed data
        return jsonify({
            'status': 'ok',
            'total_emails': len(files_by_email) - (1 if ungrouped_files else 0),
            'total_files': len(backend_data.get('files', [])),
            'total_size': total_size,
            'total_size_formatted': format_file_size(total_size),
            'files_by_email': files_by_email,
            'ungrouped_count': len(ungrouped_files),
            'sort_by': sort_by,
            'sort_order': sort_order
        }), 200
        
    except Exception as e:
        print(f"Error in get_files endpoint: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'status': 'error',
            'message': f'Internal server error: {str(e)}'
        }), 500


def format_file_size(size_bytes):
    """
    Helper function to format file size in human-readable format
    """
    if size_bytes is None or size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB", "TB"]
    import math
    i = int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return f"{s} {size_names[i]}"


def sort_files(files_list, sort_by='date', sort_order='desc'):
    """
    Helper function to sort files based on specified criteria
    
    Args:
        files_list: List of file dictionaries
        sort_by: 'name', 'size', 'date', 'type'
        sort_order: 'asc' or 'desc'
    
    Returns:
        Sorted list of files
    """
    reverse = (sort_order == 'desc')
    
    if sort_by == 'name':
        return sorted(files_list, key=lambda x: x.get('fileName', '').lower(), reverse=reverse)
    elif sort_by == 'size':
        return sorted(files_list, key=lambda x: x.get('size') or 0, reverse=reverse)
    elif sort_by == 'date':
        return sorted(
            files_list, 
            key=lambda x: x.get('uploadedAt') or x.get('lastModified') or x.get('createdOn') or '', 
            reverse=reverse
        )
    elif sort_by == 'type':
        return sorted(files_list, key=lambda x: x.get('contentType', '').lower(), reverse=reverse)
    else:
        # Default: sort by date
        return sorted(
            files_list, 
            key=lambda x: x.get('uploadedAt') or x.get('lastModified') or x.get('createdOn') or '', 
            reverse=reverse
        )



# delete the file from the bot
@app.route('/api/delete_file', methods=['POST'])
def delete_file():
    try:
        data = request.get_json()
        file_path = data.get('file_path')
        
        if not file_path:
            return jsonify({'error': 'file_path is required'}), 400
        
        # Prepare payload for backend service
        payload = {
            "file_path": file_path
        }
        
        # Call backend service with DELETE method
        backend_url = f"{BASE_URL}/delete_content_file"
        response = requests.delete(backend_url, json=payload)
        
        if response.status_code == 200:
            try:
                metadata_collection = get_mongo_connection('file-metadata')
                metadata_collection.delete_one({'blob_path': file_path})
            except Exception as exc:
                print(f"Warning: failed to remove metadata for {file_path}: {exc}")
            return jsonify({'message': 'File deleted successfully'}), 200
        else:
            return jsonify({'error': 'Failed to delete file from backend', 'details': response.text}), response.status_code
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/content/<path:file_path>', methods=['GET'])
def proxy_content_preview(file_path):
    """
    Proxy endpoint to fetch content from backend and inject CSS for white background
    This ensures we can style the content even if it's cross-origin
    """
    try:
        if 'email' not in session:
            next_url = request.full_path if request.method in ('GET', 'HEAD') else ''
            if not is_safe_url(next_url):
                next_url = ''
            return redirect(url_for('login', next=next_url or None))

        # Get theme parameter (default to 'light')
        theme = request.args.get('theme', 'light')
        
        # Construct backend URL
        backend_url = f"{BASE_URL}/content/{file_path}"
        if theme:
            separator = '&' if '?' in backend_url else '?'
            backend_url = f"{backend_url}{separator}theme={theme}"
        
        # Fetch content from backend
        response = requests.get(backend_url, timeout=30)
        content_type = response.headers.get('Content-Type', '')
        excluded_headers = {'content-encoding', 'transfer-encoding', 'connection', 'content-length'}

        # Helper to copy headers to Flask response
        def apply_headers(flask_response):
            for header, value in response.headers.items():
                header_lower = header.lower()
                if header_lower in excluded_headers:
                    continue
                if header_lower == 'content-type':
                    flask_response.headers['Content-Type'] = value
                else:
                    flask_response.headers[header] = value
            return flask_response

        # If the response is not successful, proxy it as-is (no CSS injection)
        if response.status_code != 200:
            flask_response = Response(response.content, status=response.status_code)
            return apply_headers(flask_response)
        
        # If it's HTML/text content, inject CSS for white background
        if 'text/html' in content_type or content_type.startswith('text/'):
            content = response.text
            css_injection = """
<style id="preview-override-styles">
  html, body {
    background-color: #ffffff !important;
    color: #000000 !important;
  }
  body * {
    background-color: transparent !important;
  }
  pre, code {
    background-color: #f8f9fa !important;
    color: #000000 !important;
    border: 1px solid #e9ecef !important;
    padding: 12px !important;
    border-radius: 4px !important;
  }
  .json-viewer, .json-container, [class*="json"], [id*="json"] {
    background-color: #ffffff !important;
    color: #000000 !important;
  }
  div, section, article, main {
    background-color: transparent !important;
  }
  /* Override any dark theme classes */
  [class*="dark"], [class*="black"], [class*="night"] {
    background-color: #ffffff !important;
    color: #000000 !important;
  }
</style>
"""

            if '</head>' in content:
                content = content.replace('</head>', css_injection + '</head>')
            elif '<body' in content:
                content = content.replace('<body', css_injection + '<body')
            else:
                content = css_injection + content

            flask_response = Response(content, status=200)
            return apply_headers(flask_response)

        # For binary or other content types (e.g., PDFs), proxy the content as-is
        flask_response = Response(response.content, status=200)
        return apply_headers(flask_response)
        
    except requests.exceptions.Timeout:
        return jsonify({'error': 'Backend service request timed out'}), 504
    except requests.exceptions.ConnectionError:
        return jsonify({'error': 'Could not connect to backend service'}), 503
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8000)), debug=True)